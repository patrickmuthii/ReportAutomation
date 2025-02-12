import os
from django.conf import settings
from docx import Document

def generate_report_file(report):
    """
    Generates a DOCX report file and returns the file path.
    """
    file_name = f"report_{report.id}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, file_name)

    doc = Document()
    doc.add_heading("Report Summary", level=1)
    doc.add_paragraph(f"Title: {report.title}")
    doc.add_paragraph(f"Content: {report.content}")
    doc.save(file_path)

    return file_path
